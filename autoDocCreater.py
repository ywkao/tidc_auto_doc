import os
import pandas as pd
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# import gdrive_utils as gu

class QualityControlDocGenerator:
    def __init__(self, target_folder, filename, drive='My Drive', prefix='/content/drive/'):
        self.drive = drive
        prefix = os.path.join(prefix, drive)
        self.prefix = prefix if prefix.endswith('/') else prefix + '/' # 確保 prefix 結尾有斜線
        self.base = os.path.join(self.prefix, target_folder)
        self.filename = filename
        self.csv = os.path.join(self.base, filename)
        self.doc = None
        self.black = RGBColor(0, 0, 0)
        self.blue = RGBColor(0, 0, 255)

        print(f">>> {self.base}")
        print(f">>> {self.csv}")

        # Check & Read the CSV
        self._check_folder()
        self.df = self._read_and_process_csv()

        # Find the Glue column
        # self.glue_column = self._find_column_by_keyword('Glue')
        # if not self.glue_column: print("[Warning] Glue column not found!")

    #----------------------------------------------------------------------------------------------------
    # main methods
    #----------------------------------------------------------------------------------------------------
    def create_directories(self):
        """ Create folders based on CERN ID """
        print("[INFO] 建立資料夾：")
        for index, row in self.df.iterrows():
            sub_folder = os.path.join(self.base, row['ID'])
            os.makedirs(sub_folder, exist_ok=True)
            print(f"A folder has been created: {sub_folder}")
        print("")

    def move_photos(self):
        """ Move photos to sub-directories (CERN ID) """
        print("[INFO] 移動相片：")
        for index, row in self.df.iterrows():
            self.cernID = row['ID']
            self.folder = os.path.join(self.base, row['ID'])
            os.makedirs(self.folder, exist_ok=True)

            flag, _ = self._find_path(link=row['image link'], verbosity=True)
            if flag==1: self._move_file(self.base, self.folder, row['image link'])

            flag, _ = self._find_path(link=row['p2_image link'], verbosity=True)
            if flag==1: self._move_file(self.base, self.folder, row['p2_image link'])

    def move_back_photos(self):
        """ Move back photos to sub-directories (CERN ID) """
        print("[INFO] 還原相片位置：")
        for index, row in self.df.iterrows():
            self.folder = os.path.join(self.base, row['ID'])
            os.makedirs(self.folder, exist_ok=True)

            flag, _ = self._find_path(row['image link'])
            if flag==2: self._move_file(self.folder, self.base, row['image link'])

            flag, _ = self._find_path(row['p2_image link'])
            if flag==2: self._move_file(self.folder, self.base, row['p2_image link'])

    def create_documents(self):
        """ Create documents """
        print("[INFO] 建立docx文件：")
        for index, row in self.df.iterrows():
            self._create_quality_control_doc(row)

    def move_docx(self):
        print("[INFO] 移動docx文件：")
        for index, row in self.df.iterrows():
            self.cernID = row['ID']
            self.folder = os.path.join(self.base, row['ID'])
            os.makedirs(self.folder, exist_ok=True)

            self.gdoc = row['filename+ID'] + '.docx'
            flag, _ = self._find_path(link=self.gdoc, verbosity=True)
            if flag==1: self._move_file(self.base, self.folder, self.gdoc)

    #----------------------------------------------------------------------------------------------------
    # auxiliary modules
    #----------------------------------------------------------------------------------------------------
    def _move_file(self, old, new, f):
        path1 = os.path.join(old, f)
        path2 = os.path.join(new, f)
        os.rename(path1, path2)
        print(f"- moved file from {path1} to {path2}")
        # gu.move_file(self.drive, path1, path2)

    def _create_quality_control_doc(self, row):
        self.cernID = row['ID']
        self.folder = os.path.join(self.base, self.cernID) 
        self.gdoc = row['filename+ID'] + '.docx'

        # 1st step: crreate doc at the base directory
        # 2nd step: move the doc to the target folder
        # this 2-step treatment will allow to copy links from Google Drive to excel
        self.output_file = os.path.join(self.base, self.gdoc) 

        self.doc = docx.Document()
        self._set_page_margins()
        self._add_title(row, 'title page1', 'ID')
        self._add_first_visual_inspection(row)
        self.doc.add_page_break()
        self._add_title(row, 'title page2', 'p2_ID')
        self._add_second_visual_inspection(row)
        self.doc.save(self.output_file)
        print(f"Document has been saved as {self.output_file}")

    #----------------------------------------------------------------------------------------------------
    # Load-data related
    #----------------------------------------------------------------------------------------------------
    def _check_folder(self): # TODO: return error if folder/csv does not exist
        # 確保目標資料夾存在
        if not os.path.exists(self.base):
            os.makedirs(self.base)
            print(f'新增 {self.base}')

        # List all files in the directory
        files = os.listdir(self.base)
        for f in files:
            full_path = os.path.join(self.base, f)
            file_size = os.path.getsize(full_path)
            # print(f"- {f} ({file_size} bytes)")

        # Check for specific file
        if os.path.exists(self.csv):
            print(f"[INFO] Found target file: {self.filename}")
            print(f"Full path: {self.csv}")
            print(f"File size: {os.path.getsize(self.csv)} bytes")
            print("")
        else:
            print(f"\n[ERROR] Target file does not exist: {self.filename}")

    def _read_and_process_csv(self):
        """
        Read and process the CSV file to extract relevant information
        """
        try:
            # Skip the first two rows and use the third row as headers
            df = pd.read_csv(self.csv, skiprows=2, header=0, dtype=str) # 刪除前兩行
            df = df.dropna(subset=['User'])  # 刪除 User 欄位為空的列
            df = df.fillna('') # replace all NaN with empty strings
            df.columns = [str(col).strip().replace('\n', ' ') for col in df.columns]
            # self._inspect_contents(df)
            return df

        except Exception as e:
            print(f"Error processing file: {str(e)}")
            return None

    def _inspect_contents(self, df):
        # Print the column names to verify
        print("\nColumn names:")
        for i, col in enumerate(df.columns):
            print(f"{i}: '{col}'")

        # Print the first few rows of data
        keywords = ['ID', 'Accept?', 'filename+ID', 'image link']
        print("\nFirst few rows of data (selected columns):")
        print(df[keywords].head())

    def _find_column_by_keyword(self, keyword):
        """
        Find column name that contains the given keyword
        """
        matching_cols = [col for col in self.df.columns if keyword.lower() in col.lower()]
        return matching_cols[0] if matching_cols else None

    #----------------------------------------------------------------------------------------------------
    # Document-related
    #----------------------------------------------------------------------------------------------------
    def _set_page_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_underlined_spaces(self, paragraph, Nspaces, color=None):
        full_width_underscore = '＿'
        underlined_spaces = paragraph.add_run(full_width_underscore * Nspaces)
        underlined_spaces.font.color.rgb = color or self.black
        underlined_spaces.font.underline = WD_UNDERLINE.THICK

    def _add_empty_spaces(self, paragraph, Nspaces=4, color=None):
        space = ' '
        underlined_spaces = paragraph.add_run(space * Nspaces)
        underlined_spaces.font.color.rgb = color or self.black

    def _add_title(self, row, titleKey, idKey):
        title = self.doc.add_paragraph(str(row[titleKey]))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0].font
        title_format.size = Pt(14)
        title_format.bold = True

        info = [
            ("User:"         , row['User']         if pd.notna(row['User'])         else "") ,
            ("Date:"         , row['Date']         if pd.notna(row['Date'])         else "") ,
            ("Version:"      , row['Version']      if pd.notna(row['Version'])      else "") ,
            ("Manufacturer:" , row['Manufacturer'] if pd.notna(row['Manufacturer']) else "") ,
            ("Batch:"        , row['Batch number'] if pd.notna(row['Batch number']) else "") ,
            ("ID:"           , row[idKey]          if pd.notna(row[idKey])          else "")
        ]

        for i, (key, value) in enumerate(info):
            if i % 3 == 0:
                p = self.doc.add_paragraph()

            run = p.add_run(key)
            run.font.color.rgb = self.black
            self._add_underlined_spaces(p, 2)

            run = p.add_run(value)
            run.font.underline = WD_UNDERLINE.THICK
            run.font.color.rgb = self.blue
            self._add_underlined_spaces(p, 2)

    def _process_image(self, p, image_link):
        _, image_path = self._find_path(image_link)
        if image_path is not None:
            self._add_image(p, image_path)

    def _add_formatted_table(self, row):
        # 添加新的1x2表格
        new_table = self.doc.add_table(rows=1, cols=2)
        new_table.style = 'Normal Table'
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        new_table.columns[0].width = Inches(2)
        new_table.columns[1].width = Inches(4)

        # Add chip ID & chip map
        for i, cell in enumerate(new_table.rows[0].cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i==0:
                paragraph = cell.add_paragraph()
                paragraph.add_run(f"{row['p2_Chip ID']}")
            elif i==1:
                p = cell.add_paragraph()
                self._process_image(p, row.get('p2_Chip location map link', ''))

    def _find_path(self, link, verbosity=False):
        if not link or not link.strip():
            return 0, None

        potential_paths = [
            os.path.join(self.base, link), # 1
            os.path.join(self.folder, link) # 2
        ]

        for i, path in enumerate(potential_paths):
            if os.path.exists(path):
                return i+1, path

        if verbosity:
            self._print_error(f"{link} not found for {self.cernID}")
        return 3, None

    def _print_error(self, message: str) -> None:
        print(f"\033[91m[ERROR] {message}\033[0m")

    def _add_image(self, paragraph, image_path, default_width=3.5):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(default_width))
        print(f'+ adding picture: {image_path}')

    def _add_customized_paragraph(self, paragraph, item, value, spaces):
        useEmptySpace = (spaces[0]==0) and (spaces[1]==0)

        paragraph.add_run(f"{item} ")
        self._add_underlined_spaces(paragraph, spaces[0])
        run = paragraph.add_run(value)
        if useEmptySpace is False: run.font.underline = WD_UNDERLINE.THICK
        run.font.color.rgb = self.blue
        self._add_underlined_spaces(paragraph, spaces[1])
        if useEmptySpace: self._add_empty_spaces(paragraph)

    def _add_first_visual_inspection(self, row):
        self.doc.add_heading("1st Visual Inspection – Bare PCB", level=1)

        inspection_items = [
            ("General comments:"       , row['General comments']       , 2 , [34, 0, 42]),
            ("Flatness:"               , row['Flatness']               , 1 , [2  , 2])  ,
            ("Comments:"               , row['Comments']               , 0 , [25 , 0])  ,
            ("Thickness measurements:" , row['Thickness measurements'] , 1 , [4  , 22]) ,
            ("Plating (BGA):"          , row['Plating (BGA)']          , 1 , [4  , 8])  ,
            ("Plating (Holes):"        , row['Plating (Holes)']        , 0 , [4  , 8])  ,
            ("Soldermask alignment:"   , row['Soldermask alignment']   , 1 , [4  , 26]) ,
            ("Glue problems?"          , row['Glue problems?']         , 1 , [7  , 26]) ,
            ("Test coupons (observations, continuity measurements etc.):", row['Test coupons (observations, continuity measurements etc.)'] , 2, [4, 10, 42]),
            ("Accept?"                 , row['Accept?']                , 1 , [4  , 12]),
        ]

        for item, value, nLines, spaces in inspection_items:
            if 'Accept' in item:
                p = self.doc.add_paragraph()
                self._process_image(p, row.get('image link', ''))

            if nLines > 0: p = self.doc.add_paragraph()
            self._add_customized_paragraph(p, item, value, spaces)

            if nLines == 2:
                p = self.doc.add_paragraph()
                self._add_underlined_spaces(p, spaces[2])

    def _add_second_visual_inspection(self, row):
        self.doc.add_heading("2nd Visual Inspection – Assembled PCB", level=1)

        # Assembling data
        assembled_items = [
            ("General comments:"    , row['p2_General comments']    , 1, [4, 29]),
            ("Flatness:"            , row['p2_Flatness']            , 1, [4, 30]),
            ("HGCROC type:"         , row['p2_HGCROC type']         , 1, [4,  8]),
            ("HGCROC rotation:"     , row['p2_HGCROC rotation']     , 0, [4,  8]),
            ("Connectors:"          , row['p2_Connectors']          , 1, [4,  8]),
            ("Resistors/capacitors:", row['p2_Resistors/capacitors'], 0, [4,  8]),
        ]

        for item, value, nLines, spaces in assembled_items:
            if nLines > 0: p = self.doc.add_paragraph()
            self._add_customized_paragraph(p, item, value, spaces)

        # Table for chip ID and location map
        self._add_formatted_table(row)

        # Photo at 2nd visual inspection
        p = self.doc.add_paragraph()
        self._process_image(p, row.get('p2_image link', ''))

        # Functional tests
        self.doc.add_heading("Functional Tests", level=1)
        functional_tests = [
            ("Power-on current:" , row['p2_Power-on current'] , 1, [2, 2]),
            ("Configured OK:"    , row['p2_Configured OK']    , 0, [0, 0]),
            ("Operating current:", row['p2_Operating current'], 0, [2, 2]),
            ("DAQ lines OK: "    , row['p2_DAQ lines OK']     , 1, [0, 0]),
        ]

        for item, value, nLines, spaces in functional_tests:
            if nLines > 0: p = self.doc.add_paragraph()
            self._add_customized_paragraph(p, item, value, spaces)

# Usage example
if __name__ == "__main__":
    folder, filename = 'autoDoc', 'test.csv'
    generator = QualityControlDocGenerator(folder, filename)
    generator.create_directories()
    generator.create_documents()
